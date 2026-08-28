# -*- coding: utf-8 -*-
"""Compone riassunto.docx dai file riassunti/capNN.md."""
import glob, os, re, json, hashlib
import schemi
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK   = RGBColor(0x1c, 0x27, 0x33)
BLUE  = RGBColor(0x2c, 0x4a, 0x6b)
GREY  = RGBColor(0x5b, 0x6b, 0x7c)
IMG   = "build/img"
os.makedirs(IMG, exist_ok=True)

TITOLO = "Il controllo di gestione"
SOTTO  = "Busco · Giovannoni · Riccaboni — V edizione, Wolters Kluwer"

PARTI = {
    1: "Parte I — Fondamenti", 4: "Parte II — Metodologie, strumenti ed esperienze",
    16: "Parte III — Pratiche innovative", 30: "Parte IV — Casi ed esperienze",
}

# ---------------------------------------------------------------- utilità docx
def _fld(par, instr):
    r = par.add_run()
    for t, txt in (("begin", None), ("instrText", instr), ("separate", None),
                   ("t", "aggiorna con F9"), ("end", None)):
        e = OxmlElement("w:" + ("fldChar" if t in ("begin", "separate", "end") else t))
        if t in ("begin", "separate", "end"):
            e.set(qn("w:fldCharType"), t)
        else:
            e.set(qn("xml:space"), "preserve"); e.text = txt
        r._r.append(e)

def _shade(cell, hexcol):
    el = OxmlElement("w:shd"); el.set(qn("w:val"), "clear"); el.set(qn("w:fill"), hexcol)
    cell._tc.get_or_add_tcPr().append(el)

def _page_numbers(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _fld(p, "PAGE")
    for r in p.runs:
        r.font.size = Pt(8.5); r.font.color.rgb = GREY

def _rich(par, text):
    """Rende **grassetto**, *corsivo*, `code` dentro un paragrafo."""
    for tok in re.split(r"(\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = par.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        else:
            par.add_run(tok)

def _styles(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(10.5); n.font.color.rgb = INK
    n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.14
    for nome, sz, col, before, after in (("Heading 1", 19, BLUE, 22, 10),
                                         ("Heading 2", 14, BLUE, 16, 6),
                                         ("Heading 3", 11.5, INK, 12, 4)):
        st = doc.styles[nome]
        st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.bold = True
        st.font.color.rgb = col
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

# ---------------------------------------------------------------- tabelle
def tabella_compare(doc, s):
    cols, rows = s["columns"], s["rows"]
    ncol = len(rows[0])
    head = cols if len(cols) == ncol else [""] + cols
    t = doc.add_table(rows=1, cols=ncol); t.style = "Table Grid"
    t.autofit = True
    for i, h in enumerate(head):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; r = p.add_run(h); r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = BLUE
        _shade(c, "EDF2F7")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            _rich(p, str(v))
            for r in p.runs:
                r.font.size = Pt(9)
            if i == 0 and len(cols) != ncol:
                for r in p.runs: r.bold = True
    doc.add_paragraph()

def didascalia(doc, testo):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(testo); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(12)

# ---------------------------------------------------------------- markdown
def scrivi_markdown(doc, md, nfig):
    righe = md.split("\n")
    i = 0
    while i < len(righe):
        ln = righe[i]; i += 1
        s = ln.strip()
        if not s:
            continue
        if s.startswith("# "):
            continue                       # il titolo è già stato scritto
        if s.startswith("## "):
            doc.add_heading(s[3:], level=2); continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3); continue
        if s.startswith("|"):              # tabella markdown
            blocco = [s]
            while i < len(righe) and righe[i].strip().startswith("|"):
                blocco.append(righe[i].strip()); i += 1
            tabella_markdown(doc, blocco); continue
        if s.startswith("> "):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
            _rich(p, s[2:])
            for r in p.runs:
                r.font.size = Pt(9.8); r.font.color.rgb = BLUE
            continue
        if re.match(r"^[-*] ", s):
            p = doc.add_paragraph(style="List Bullet"); _rich(p, s[2:]); continue
        if re.match(r"^\d+\. ", s):
            p = doc.add_paragraph(style="List Number")
            _rich(p, re.sub(r"^\d+\. ", "", s)); continue
        if set(s) <= set("-—_") and len(s) > 2:
            continue
        p = doc.add_paragraph(); _rich(p, s)

def tabella_markdown(doc, blocco):
    righe = [r for r in blocco if not re.match(r"^\|[\s:|-]+\|$", r)]
    dati = [[c.strip() for c in r.strip("|").split("|")] for r in righe]
    if not dati: return
    t = doc.add_table(rows=1, cols=len(dati[0])); t.style = "Table Grid"
    for i, h in enumerate(dati[0]):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = BLUE; _shade(c, "EDF2F7")
    for row in dati[1:]:
        cells = t.add_row().cells
        for i, v in enumerate(row[:len(dati[0])]):
            cells[i].text = ""
            _rich(cells[i].paragraphs[0], v)
            for r in cells[i].paragraphs[0].runs: r.font.size = Pt(9)
    doc.add_paragraph()

# ---------------------------------------------------------------- documento
def costruisci(out="riassunto.docx"):
    doc = Document()
    _styles(doc)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.4)

    # --- frontespizio
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RIASSUNTO RAGIONATO"); r.font.size = Pt(12); r.font.color.rgb = GREY
    r.bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITOLO); r.font.size = Pt(34); r.bold = True; r.font.color.rgb = BLUE
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SOTTO); r.font.size = Pt(12); r.font.color.rgb = GREY
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sintesi per capitoli, con schemi e mappe concettuali")
    r.font.size = Pt(11); r.italic = True; r.font.color.rgb = GREY
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- indice
    doc.add_heading("Indice", level=1)
    p = doc.add_paragraph()
    _fld(p, r'TOC \o "1-2" \h \z \u')
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    _page_numbers(sec)

    nfig = 0
    files = sorted(glob.glob("riassunti/cap*.md"))
    for f in files:
        num = int(re.search(r"cap(\d+)", f).group(1))
        md = open(f).read()
        titolo = md.split("\n")[0].lstrip("# ").strip()

        if num in PARTI:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for _ in range(4): doc.add_paragraph()
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(PARTI[num]); r.font.size = Pt(22); r.bold = True
            r.font.color.rgb = BLUE
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        doc.add_heading(titolo, level=1)

        schemi_cap = schemi.parse_schemi(md)
        corpo = schemi.strip_schemi(md)
        # rimuove la riga del titolo
        corpo = "\n".join(corpo.split("\n")[1:])
        scrivi_markdown(doc, corpo, nfig)

        if schemi_cap:
            doc.add_heading("Schemi di sintesi", level=2)
            for k, sc in enumerate(schemi_cap):
                if sc["type"] == "compare":
                    p = doc.add_paragraph()
                    r = p.add_run(sc["title"]); r.bold = True; r.font.size = Pt(10)
                    r.font.color.rgb = BLUE
                    tabella_compare(doc, sc)
                    if sc.get("note"): didascalia(doc, sc["note"])
                else:
                    nfig += 1
                    path = f"{IMG}/cap{num:02d}_{k}.png"
                    if schemi.render(sc, path):
                        doc.add_picture(path, width=Cm(16.2))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        didascalia(doc, f"Schema {nfig} — {sc['title']}")
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.save(out)
    return out, len(files), nfig

if __name__ == "__main__":
    out, n, f = costruisci()
    print(f"{out}: {n} capitoli, {f} schemi")
